#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

BASE = "/Users/zeno/Desktop/CLAUDE_WORKSPACE_PERSONALE/IL LENZUOLO BIANCO/"
OUTPUT_DOCX = BASE + "REEL_TEASER_1.docx"
OUTPUT_PDF  = BASE + "REEL_TEASER_1.pdf"

DATA_EVENTO = "mercoledì 11 giugno 2026"

SEQUENZA = [
    ("1", "0–4 sec",
     "Mani che aprono lentamente un grande lenzuolo bianco — slow motion — sfondo neutro o prato verde",
     "(nessun testo)"),
    ("2", "4–8 sec",
     "Il lenzuolo che si posa sul prato, vento leggero che lo muove",
     "(nessun testo)"),
    ("3", "8–13 sec",
     "Le mura di Ferrara al tramonto — luce calda, cielo che vira all'arancio",
     "(nessun testo)"),
    ("4", "13–17 sec",
     "Schermata nera o dissolvenza lenta",
     "«mercoledì 11 giugno»"),
    ("5", "17–20 sec",
     "Schermata nera",
     "«ore 19.00 — sottomura»"),
    ("6", "20–25 sec",
     "Logo / testo finale centrato su sfondo scuro",
     "«il lenzuolo bianco» + «#illenznuolobianco»"),
]

INDICAZIONI = [
    "Girato con iPhone in modalità cinematica o slow motion 240fps per la scena del lenzuolo",
    "Colori: non filtrare troppo — la luce naturale del tramonto è già perfetta",
    "Niente stabilizzatore elettronico sulla scena del lenzuolo — il movimento naturale è più vero",
    "Testo sovrapposto: font serif sottile, bianco, centrato — niente grassetto, niente ombra",
    "Esportare a 1080×1920 (9:16), 30fps",
]

AUDIO = [
    ("Opzione A — Ambient piano",
     "Nils Frahm / Ólafur Arnalds — piano ambient lento, 60–80 BPM. Entra al taglio 1, volume basso."),
    ("Opzione B — Silenzio + fade",
     "Silenzio con solo il suono del vento nei primi 13 secondi, poi musica in fade al taglio 4."),
    ("Opzione C — Solo sound design",
     "Fruscio del lenzuolo, vento, uccelli in lontananza — zero musica. Scelta più coraggiosa, molto efficace su TikTok."),
]

CAPTION = (
    "c'è un'idea che gira da un po'.\n"
    "ha bisogno di un prato, di un lenzuolo e di una serata di giugno.\n\n"
    f"{DATA_EVENTO} — ore 19.00\n"
    "sottomura di Ferrara\n\n"
    "#illenznuolobianco"
)

# ─── HELPERS DOCX ─────────────────────────────────────────────────────────────

def set_font(run, name="Georgia", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sizes = {1: 22, 2: 15, 3: 12}
    colors_map = {1: (20, 20, 20), 2: (50, 50, 50), 3: (80, 80, 80)}
    set_font(r, "Georgia", sizes.get(level, 12), bold=True, color=colors_map.get(level))
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, italic=False, color=(60, 60, 60), indent_cm=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, "Georgia", 11, italic=italic, color=color)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, "Georgia", 11, color=(60, 60, 60))
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)

def shade_cell(cell, fill="D8D8D8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

# ─── DOCX ─────────────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # Copertina
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("IL LENZUOLO BIANCO")
    set_font(r, "Georgia", 26, bold=True, color=(20, 20, 20))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Script — Reel Teaser #1")
    set_font(r2, "Georgia", 14, italic=True, color=(80, 80, 80))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Ferrara  •  {DATA_EVENTO}  •  {datetime.date.today().strftime('%d/%m/%Y')}")
    set_font(r3, "Georgia", 10, color=(140, 140, 140))
    p3.paragraph_format.space_before = Pt(6)

    doc.add_paragraph()

    # Scheda tecnica
    add_heading(doc, "SCHEDA TECNICA", 1)
    specs = [
        ("Formato",   "Reel verticale 9:16"),
        ("Durata",    "20–25 secondi"),
        ("Tono",      "Lento, evocativo — nessuna voce fuori campo"),
        ("Obiettivo", "Curiosità e FOMO — non spiegare, suggerire"),
        ("Pubblico",  "Ferraresi e chi conosce le mura — poi espandere"),
    ]
    table = doc.add_table(rows=len(specs), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(specs):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                run.bold = True
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
        if i % 2 == 0:
            shade_cell(row.cells[0], "EEEEEE")
            shade_cell(row.cells[1], "EEEEEE")

    doc.add_paragraph()

    # Sequenza visiva
    add_heading(doc, "SEQUENZA VISIVA", 1)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    hdrs = ["#", "Timing", "Immagine", "Testo sovrapposto"]
    for i, (cell, h) in enumerate(zip(tbl2.rows[0].cells, hdrs)):
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                run.bold = True
        shade_cell(cell, "D0D0D0")

    for num, timing, img, testo in SEQUENZA:
        row = tbl2.add_row()
        data = [num, timing, img, testo]
        for i, (cell, text) in enumerate(zip(row.cells, data)):
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True

    doc.add_paragraph()

    # Indicazioni tecniche
    add_heading(doc, "INDICAZIONI TECNICHE", 1)
    for ind in INDICAZIONI:
        add_bullet(doc, ind)

    doc.add_paragraph()

    # Audio
    add_heading(doc, "AUDIO", 1)
    for title, desc in AUDIO:
        p = doc.add_paragraph()
        r1 = p.add_run(title + "  ")
        set_font(r1, "Georgia", 11, bold=True, color=(40, 40, 40))
        r2 = p.add_run(desc)
        set_font(r2, "Georgia", 11, color=(70, 70, 70))
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    # Caption
    add_heading(doc, "CAPTION DEL POST", 1)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    r = p.add_run(CAPTION)
    set_font(r, "Georgia", 11, italic=True, color=(60, 60, 60))
    p.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT_DOCX)
    print(f"✓ DOCX salvato: {OUTPUT_DOCX}")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def build_pdf():
    doc = SimpleDocTemplate(
        OUTPUT_PDF, pagesize=A4,
        leftMargin=3*cm, rightMargin=3*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    styles = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    cover_title = sty("CT", fontSize=26, leading=32, alignment=TA_CENTER,
                      fontName="Times-Bold", textColor=colors.HexColor("#141414"), spaceAfter=8)
    cover_sub   = sty("CS", fontSize=14, leading=20, alignment=TA_CENTER,
                      fontName="Times-Italic", textColor=colors.HexColor("#505050"), spaceAfter=6)
    cover_meta  = sty("CM", fontSize=10, leading=14, alignment=TA_CENTER,
                      fontName="Times-Roman", textColor=colors.HexColor("#8C8C8C"))
    h1 = sty("H1", fontSize=16, leading=22, fontName="Times-Bold",
             textColor=colors.HexColor("#141414"), spaceBefore=18, spaceAfter=8)
    body = sty("B", fontSize=11, leading=16, fontName="Times-Roman",
               textColor=colors.HexColor("#3C3C3C"), spaceAfter=5, alignment=TA_JUSTIFY)
    italic_b = sty("IB", fontSize=11, leading=16, fontName="Times-Italic",
                   textColor=colors.HexColor("#505050"), spaceAfter=8, leftIndent=30,
                   rightIndent=20, alignment=TA_LEFT)
    bullet_s = sty("BL", fontSize=11, leading=16, fontName="Times-Roman",
                   textColor=colors.HexColor("#3C3C3C"), leftIndent=20, spaceAfter=4)

    story = []

    def sp(n=8):
        story.append(Spacer(1, n))

    def hr():
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#C0C0C0"), spaceAfter=8))

    # Copertina
    story += [
        Spacer(1, 3.5*cm),
        Paragraph("IL LENZUOLO BIANCO", cover_title),
        Paragraph("Script — Reel Teaser #1", cover_sub),
        Paragraph(f"Ferrara  •  {DATA_EVENTO}  •  {datetime.date.today().strftime('%d/%m/%Y')}",
                  cover_meta),
        Spacer(1, 1.5*cm),
    ]

    # Scheda tecnica
    story.append(Paragraph("SCHEDA TECNICA", h1))
    hr()
    specs = [
        ["Formato",   "Reel verticale 9:16"],
        ["Durata",    "20–25 secondi"],
        ["Tono",      "Lento, evocativo — nessuna voce fuori campo"],
        ["Obiettivo", "Curiosità e FOMO — non spiegare, suggerire"],
        ["Pubblico",  "Ferraresi e chi conosce le mura — poi espandere"],
    ]
    tbl_spec = Table(specs, colWidths=[4*cm, 13.5*cm])
    tbl_spec.setStyle(TableStyle([
        ("FONTNAME",    (0,0),(0,-1),  "Times-Bold"),
        ("FONTNAME",    (1,0),(1,-1),  "Times-Roman"),
        ("FONTSIZE",    (0,0),(-1,-1), 11),
        ("LEADING",     (0,0),(-1,-1), 15),
        ("ROWBACKGROUNDS",(0,0),(-1,-1),[colors.HexColor("#F5F5F5"), colors.white]),
        ("GRID",        (0,0),(-1,-1), 0.4, colors.HexColor("#C8C8C8")),
        ("TOPPADDING",  (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING", (0,0),(-1,-1), 8),
    ]))
    story.append(tbl_spec)
    sp()

    # Sequenza visiva
    story.append(Paragraph("SEQUENZA VISIVA", h1))
    hr()
    seq_data = [["#", "Timing", "Immagine", "Testo"]]
    for num, timing, img, testo in SEQUENZA:
        seq_data.append([num, timing, img, testo])

    tbl_seq = Table(seq_data, colWidths=[0.7*cm, 2*cm, 10.3*cm, 4.5*cm], repeatRows=1)
    tbl_seq.setStyle(TableStyle([
        ("FONTNAME",    (0,0),(-1,0),  "Times-Bold"),
        ("FONTNAME",    (0,1),(-1,-1), "Times-Roman"),
        ("FONTNAME",    (0,1),(0,-1),  "Times-Bold"),
        ("FONTSIZE",    (0,0),(-1,-1), 10),
        ("LEADING",     (0,0),(-1,-1), 14),
        ("BACKGROUND",  (0,0),(-1,0),  colors.HexColor("#D0D0D0")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#F5F5F5")]),
        ("GRID",        (0,0),(-1,-1), 0.4, colors.HexColor("#C0C0C0")),
        ("VALIGN",      (0,0),(-1,-1), "TOP"),
        ("TOPPADDING",  (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING", (0,0),(-1,-1), 5),
        ("RIGHTPADDING",(0,0),(-1,-1), 5),
    ]))
    story.append(tbl_seq)
    sp()

    # Indicazioni tecniche
    story.append(Paragraph("INDICAZIONI TECNICHE", h1))
    hr()
    for ind in INDICAZIONI:
        story.append(Paragraph(f"• {ind}", bullet_s))
    sp()

    # Audio
    story.append(Paragraph("AUDIO", h1))
    hr()
    for title, desc in AUDIO:
        story.append(Paragraph(
            f"<b>{title}</b><br/>{desc}",
            sty(f"AUD{title[:3]}", fontSize=11, leading=16, fontName="Times-Roman",
                textColor=colors.HexColor("#3C3C3C"), leftIndent=10, spaceAfter=10,
                alignment=TA_LEFT)
        ))
    sp()

    # Caption
    story.append(Paragraph("CAPTION DEL POST", h1))
    hr()
    story.append(Paragraph(
        CAPTION.replace("\n", "<br/>"),
        sty("CAP", fontSize=11, leading=17, fontName="Times-Italic",
            textColor=colors.HexColor("#505050"), leftIndent=30, rightIndent=20,
            spaceAfter=10, alignment=TA_LEFT)
    ))
    sp(20)
    story.append(Paragraph(
        "Il Lenzuolo Bianco — Ferrara",
        sty("FTR", fontSize=9, fontName="Times-Italic",
            alignment=TA_CENTER, textColor=colors.HexColor("#AAAAAA"))
    ))

    doc.build(story)
    print(f"✓ PDF salvato:  {OUTPUT_PDF}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("\nFatto!")
