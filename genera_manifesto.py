#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER

BASE = "/Users/zeno/Desktop/CLAUDE_WORKSPACE_PERSONALE/IL LENZUOLO BIANCO/"
OUTPUT_DOCX = BASE + "MANIFESTO_STAMPABILE.docx"
OUTPUT_PDF  = BASE + "MANIFESTO_STAMPABILE.pdf"

TITOLO    = "IL LENZUOLO BIANCO"
SOTTOTIT  = "Ferrara — sottomura — mercoledì 11 giugno 2026 — ore 19.00"
HASHTAG   = "#illenznuolobianco"

CORPO = [
    ("Un mercoledì di giugno, all'imbrunire.", False),
    ("Sottomura di Ferrara, sul prato.", False),
    ("", False),
    ("Porta un lenzuolo bianco —", False),
    ("grande, vecchio, meglio se di canapa.", False),
    ("Porta qualcosa da bere, freddo.", False),
    ("Porta qualcosa da mangiare, buono.", False),
    ("Porta la tua musica. Non troppo alta.", False),
    ("", False),
    ("Stenditi sul prato.", False),
    ("Conosci il tuo vicino di lenzuolo.", False),
    ("Scambia un sorso, un morso, una canzone.", False),
    ("", False),
    ("Non c'è un programma.", False),
    ("Non c'è un biglietto.", False),
    ("Non c'è un dress code.", False),
    ("Non c'è un palco.", False),
    ("", False),
    ("C'è solo il prato, le mura,", False),
    ("il cielo che cambia colore", False),
    ("e la voglia di stare bene insieme —", False),
    ("semplicemente.", False),
    ("", False),
    ("Prima di andare,", False),
    ("lascia il posto come lo hai trovato.", False),
    ("", False),
    ("Il resto, inventalo.", True),   # True = grassetto
]

# ─── DOCX ─────────────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(3.5)
        section.bottom_margin = Cm(3.5)
        section.left_margin   = Cm(4)
        section.right_margin  = Cm(4)

    def para(text, size=12, bold=False, italic=False,
             color=(20, 20, 20), space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.font.name  = "Georgia"
            r.font.size  = Pt(size)
            r.bold       = bold
            r.italic     = italic
            r.font.color.rgb = RGBColor(*color)
        return p

    # Titolo principale
    para(TITOLO, size=30, bold=True, space_before=0, space_after=6, color=(10, 10, 10))

    # Linea decorativa — paragrafo vuoto con bordo inferiore tramite XML
    p_hr = doc.add_paragraph()
    p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hr.paragraph_format.space_before = Pt(2)
    p_hr.paragraph_format.space_after  = Pt(14)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Corpo del manifesto
    for line, bold in CORPO:
        if line == "":
            para("", space_before=0, space_after=8)
        else:
            size   = 13 if bold else 12
            color  = (10, 10, 10) if bold else (40, 40, 40)
            para(line, size=size, bold=bold, italic=not bold,
                 color=color, space_before=0, space_after=5)

    # Spazio prima del footer
    para("", space_before=0, space_after=16)

    # Linea decorativa inferiore
    p_hr2 = doc.add_paragraph()
    p_hr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hr2.paragraph_format.space_before = Pt(2)
    p_hr2.paragraph_format.space_after  = Pt(10)
    pPr2 = p_hr2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"),   "single")
    top2.set(qn("w:sz"),    "4")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "AAAAAA")
    pBdr2.append(top2)
    pPr2.append(pBdr2)

    # Sottotitolo / dettagli evento
    para(SOTTOTIT, size=10, italic=True, color=(100, 100, 100), space_after=4)
    para(HASHTAG,  size=10, bold=True,  color=(80, 80, 80),   space_after=0)

    doc.save(OUTPUT_DOCX)
    print(f"✓ DOCX salvato: {OUTPUT_DOCX}")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def build_pdf():
    doc = SimpleDocTemplate(
        OUTPUT_PDF, pagesize=A4,
        leftMargin=4*cm, rightMargin=4*cm,
        topMargin=3.5*cm, bottomMargin=3.5*cm,
    )

    def sty(name, **kw):
        return ParagraphStyle(name, **kw)

    DARK    = colors.HexColor("#0A0A0A")
    MID     = colors.HexColor("#282828")
    BODY_C  = colors.HexColor("#282828")
    GREY    = colors.HexColor("#646464")
    LGREY   = colors.HexColor("#AAAAAA")

    title_sty = sty("TIT", fontSize=32, leading=38, fontName="Times-Bold",
                    textColor=DARK, alignment=TA_CENTER, spaceAfter=4)
    sub_sty   = sty("SUB", fontSize=10, leading=14, fontName="Times-Italic",
                    textColor=GREY, alignment=TA_CENTER, spaceAfter=3)
    hash_sty  = sty("HSH", fontSize=10, leading=14, fontName="Times-Bold",
                    textColor=colors.HexColor("#505050"), alignment=TA_CENTER, spaceAfter=0)
    line_sty  = sty("LN",  fontSize=12, leading=18, fontName="Times-Italic",
                    textColor=BODY_C, alignment=TA_CENTER, spaceAfter=4)
    bold_line = sty("BL",  fontSize=13, leading=20, fontName="Times-Bold",
                    textColor=MID, alignment=TA_CENTER, spaceAfter=4)
    empty_sty = sty("EM",  fontSize=6,  leading=10, fontName="Times-Roman",
                    textColor=colors.white, alignment=TA_CENTER)

    def hr(thick=0.5, color=LGREY, before=6, after=12):
        return HRFlowable(width="80%", thickness=thick, color=color,
                          spaceAfter=after, spaceBefore=before,
                          hAlign="CENTER")

    story = []

    # Titolo
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(TITOLO, title_sty))
    story.append(hr(thick=0.6, before=2, after=16))

    # Corpo
    for line, bold in CORPO:
        if line == "":
            story.append(Paragraph(" ", empty_sty))
        elif bold:
            story.append(Paragraph(line, bold_line))
        else:
            story.append(Paragraph(line, line_sty))

    story.append(Spacer(1, 0.6*cm))
    story.append(hr(thick=0.6, before=2, after=10))

    # Footer
    story.append(Paragraph(SOTTOTIT, sub_sty))
    story.append(Spacer(1, 0.15*cm))
    story.append(Paragraph(HASHTAG, hash_sty))

    doc.build(story)
    print(f"✓ PDF salvato:  {OUTPUT_PDF}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("\nFatto!")
