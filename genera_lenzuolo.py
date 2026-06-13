#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_DOCX = "/Users/zeno/Desktop/IL_LENZUOLO_BIANCO.docx"
OUTPUT_PDF  = "/Users/zeno/Desktop/IL_LENZUOLO_BIANCO.pdf"

# ─── contenuto ────────────────────────────────────────────────────────────────

MANIFESTO = [
    "Il Lenzuolo Bianco",
    "",
    "Un mercoledì di giugno, all'imbrunire.",
    "Sottomura di Ferrara, sul prato.",
    "",
    "Porta un lenzuolo bianco — grande, vecchio, meglio se di canapa.",
    "Porta qualcosa da bere, freddo.",
    "Porta qualcosa da mangiare, buono.",
    "Porta la tua musica. Non troppo alta.",
    "",
    "Stenditi sul prato.",
    "Conosci il tuo vicino di lenzuolo.",
    "Scambia un sorso, un morso, una canzone.",
    "",
    "Non c'è un programma.",
    "Non c'è un biglietto.",
    "Non c'è un dress code.",
    "Non c'è un palco.",
    "",
    "C'è solo il prato, le mura, il cielo che cambia colore",
    "e la voglia di stare bene insieme — semplicemente.",
    "",
    "Prima di andare, lascia il posto come lo hai trovato.",
    "",
    "Il resto, inventalo.",
    "",
    "— Il Lenzuolo Bianco, Ferrara",
]

PIANO_EDITORIALE = [
    ("mer 20 mag", "IG story", "Primo teaser — solo frase misteriosa", "Curiosità"),
    ("ven 22 mag", "IG post", "Foto mura al tramonto + frase minima", "Atmosfera"),
    ("lun 25 mag", "IG story", "Foto lenzuolo di canapa + «lo stai cercando?»", "Engagement"),
    ("mer 27 mag", "IG post", "Annuncio ufficiale completo + manifesto", "Diffusione"),
    ("ven 29 mag", "IG story", "Sondaggio interattivo «cosa porti?»", "Partecipazione"),
    ("dom 31 mag", "IG Reel", "Teaser visivo 15 sec", "Reach organico"),
    ("lun 1 giu", "IG post", "«Guida minima al lenzuolo bianco»", "Istruzioni chiare"),
    ("mer 3 giu", "IG story", "Countdown — «mancano 8 giorni»", "Urgenza dolce"),
    ("ven 5 giu", "IG post", "«Dove trovare il lenzuolo» — mercatini, soffitte", "Racconto + pratico"),
    ("dom 7 giu", "IG story", "Condividi chi hai già convinto a venire", "Social proof"),
    ("lun 9 giu", "IG Reel", "Secondo teaser, più concreto — le mura, il prato", "Spinta finale"),
    ("mar 10 giu", "IG story", "«Domani. Ore 19. Sottomura.» — nient'altro", "Tensione"),
    ("mer 11 giu", "IG story live", "Durante l'evento — 2/3 storie max", "Autenticità"),
    ("gio 12 giu", "IG post", "La foto più bella dell'evento", "Emozione, FOMO"),
    ("ven 13 giu", "IG gallery", "Serie di scatti — lenzuoli, persone, dettagli", "Racconto"),
    ("dom 15 giu", "IG Reel", "Video riassuntivo emozionale", "Viralità posticipata"),
    ("lun 16 giu", "IG post", "Annuncio seconda data", "Conversione"),
]

CRESCITA = [
    ("1ª — giugno 2026", "Creare il materiale visivo perfetto", "30–50 persone selezionate"),
    ("2ª — luglio o settembre 2026", "Passaparola + chi non c'era vuole esserci", "Apri a chiunque, annuncio social"),
    ("3ª+", "Evento autonomo, la gente si organizza da sola", "Community, hashtag, edizioni in altre città"),
]

# ─── DOCX ─────────────────────────────────────────────────────────────────────

def set_font(run, name="Georgia", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(30, 30, 30)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13}
    set_font(run, "Georgia", sizes.get(level, 12), bold=True, color=color)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, italic=False, indent=False, color=(60, 60, 60)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Georgia", 11, italic=italic, color=color)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, indent_cm=1):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "Georgia", 11, color=(60, 60, 60))
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(3)

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                if i == 0 and bold_first:
                    run.bold = True

def shade_cell(cell, fill="E8E8E8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def build_docx():
    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ── Copertina ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("IL LENZUOLO BIANCO")
    set_font(r, "Georgia", 28, bold=True, color=(20, 20, 20))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Piano evento virale — Ferrara, giugno 2026")
    set_font(r2, "Georgia", 13, italic=True, color=(100, 100, 100))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Ideato da Zeno Govoni  •  {datetime.date.today().strftime('%d/%m/%Y')}")
    set_font(r3, "Georgia", 10, color=(140, 140, 140))
    p3.paragraph_format.space_before = Pt(8)

    doc.add_page_break()

    # ── 1. Il concept ──
    add_heading(doc, "1. IL CONCEPT", 1)
    add_body(doc, (
        "Il Lenzuolo Bianco è un evento flash mob ad ingresso libero, "
        "che si tiene il mercoledì sera all'imbrunire, a giugno, nel sottomura di Ferrara. "
        "I partecipanti portano un grande lenzuolo bianco — meglio se di canapa, trovato in un mercatino — "
        "si stendono sul prato esterno alle mura, portando cibo, bevande fredde e la propria musica. "
        "L'evento si conclude lasciando il luogo esattamente come era all'arrivo."
    ))

    add_heading(doc, "Perché può diventare virale", 2)
    bullets = [
        "Visivo e fotografabile — lenzuola bianche sul verde, mura medievali, luce del tramonto",
        "Bassa barriera d'ingresso — nessun biglietto, nessuna prenotazione, nessun dress code",
        "Altamente condivisibile — ogni lenzuolo è un protagonista diverso",
        "Nostalgia + creatività — il lenzuolo di canapa è un oggetto con anima",
        "La regola del rispetto — «lascia il posto come lo hai trovato» è un valore, non solo una norma",
    ]
    for b in bullets:
        add_bullet(doc, b)

    # ── 2. Prima edizione ──
    add_heading(doc, "2. PRIMA EDIZIONE — PICCOLA E CURATA", 1)
    add_body(doc, "Non fare grande subito. Fai bello.", italic=True)
    steps = [
        "Invita 30–50 persone che conosci e che sai che postano bene",
        "Scegli un mercoledì di giugno con meteo affidabile (metà mese)",
        "Scatta foto e video con un fotografo amico, non uno studio pagato",
        "Crea contenuti prima, durante e dopo l'evento",
        "L'obiettivo del primo evento non è la folla — è il materiale visivo per le edizioni successive",
    ]
    for s in steps:
        add_bullet(doc, s)

    # ── 3. Identità visiva ──
    add_heading(doc, "3. IDENTITÀ VISIVA E NOME", 1)
    add_body(doc, (
        "Il nome «Il Lenzuolo Bianco» è evocativo, italiano, ironico, caldo — "
        "non va cambiato. Va costruita un'identità visiva coerente attorno ad esso."
    ))
    items = [
        "Logo: minimalista — un lenzuolo che sventola o una sagoma stilizzata",
        "Palette: bianco / verde prato / mattone rosso delle mura ferraresi",
        "Hashtag: #illenznuolobianco — uno solo, usato sempre",
        "Profilo Instagram dedicato — separato da Street Dinner e Hotel Annunziata",
    ]
    for i in items:
        add_bullet(doc, i)

    # ── 4. Il manifesto ──
    add_heading(doc, "4. IL MANIFESTO", 1)
    add_body(doc, "Da stampare, condividere, incorniciare.", italic=True)
    doc.add_paragraph()

    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.LEFT
    box.paragraph_format.left_indent  = Cm(1.5)
    box.paragraph_format.right_indent = Cm(1.5)
    box.paragraph_format.space_before = Pt(6)
    box.paragraph_format.space_after  = Pt(6)

    for line in MANIFESTO:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent  = Cm(2)
        p.paragraph_format.right_indent = Cm(2)
        p.paragraph_format.space_after  = Pt(2)
        is_title = line == "Il Lenzuolo Bianco"
        is_firma = line.startswith("—")
        r = p.add_run(line)
        if is_title:
            set_font(r, "Georgia", 14, bold=True, color=(20, 20, 20))
        elif is_firma:
            set_font(r, "Georgia", 10, italic=True, color=(120, 120, 120))
        else:
            set_font(r, "Georgia", 11, italic=True, color=(50, 50, 50))

    # ── 5. Comunicazione social ──
    add_heading(doc, "5. COMUNICAZIONE SOCIAL", 1)

    add_heading(doc, "Tono di voce", 2)
    add_body(doc, (
        "Non è un evento. È un'idea. Il tono è sempre personale, mai da «agenzia eventi». "
        "Breve, poetico, con un pizzico di ironia. Niente maiuscole urlate, niente emoji a raffica."
    ))

    add_heading(doc, "Bio Instagram suggerita", 2)
    for line in [
        "il lenzuolo bianco",
        "un prato. le mura. un mercoledì di giugno.",
        "📍 Ferrara — sottomura",
        "#illenznuolobianco",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        r = p.add_run(line)
        set_font(r, "Courier New", 10, color=(50, 50, 50))
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "Fase 1 — Teaser (3 settimane prima)", 2)
    captions_teaser = [
        ("Post 1 — foto mura al tramonto",
         "c'è un'idea che gira da un po'.\nha bisogno di un prato, di qualche lenzuolo e di una serata di giugno.\npresto."),
        ("Post 2 — mano con lenzuolo di canapa",
         "nei mercatini ne trovi ancora.\npesanti, ruvidi, bianchi.\nperfetti."),
        ("Post 3 — prato vuoto, mura sullo sfondo",
         "mercoledì 11 giugno.\nore 19.00.\nsottomura, lato est.\nporta un lenzuolo."),
    ]
    for title, caption in captions_teaser:
        add_body(doc, title, italic=False, color=(80, 80, 80))
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        r = p.add_run(caption)
        set_font(r, "Georgia", 10, italic=True, color=(90, 90, 90))
        p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Fase 2 — Annuncio ufficiale (2 settimane prima)", 2)
    annuncio = (
        "Il Lenzuolo Bianco\n\n"
        "un evento senza palco, senza biglietto, senza programma.\n\n"
        "porti un lenzuolo bianco — grande, meglio se vecchio.\n"
        "porti da bere, da mangiare, la tua musica.\n"
        "stenditi sul prato delle mura.\n"
        "conosci chi ti sta accanto.\n\n"
        "mercoledì 11 giugno, ore 19.00\n"
        "sottomura di Ferrara — prato esterno lato Rampari\n\n"
        "prima di andare, lasci il posto come l'hai trovato.\n\n"
        "#illenznuolobianco"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(annuncio)
    set_font(r, "Georgia", 10, italic=True, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Fase 3 — Hype (settimana prima)", 2)
    guida = (
        "guida minima al lenzuolo bianco:\n\n"
        "✦ il lenzuolo: bianco. più vecchio meglio è. di canapa se ce l'hai.\n"
        "✦ da bere: freddo. buono. condivisibile.\n"
        "✦ da mangiare: niente piatti caldi. aperitivo lungo.\n"
        "✦ musica: sì, ma con rispetto del vicino di lenzuolo.\n"
        "✦ orario: dalle 19. finché c'è luce e voglia.\n"
        "✦ il prato: lascialo come lo trovi.\n\n"
        "il resto è tuo."
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(guida)
    set_font(r, "Georgia", 10, italic=True, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Fase 4 — Post evento", 2)
    post_evento = (
        "eravate in tanti.\n"
        "ognuno con il suo lenzuolo, il suo cibo, la sua musica.\n"
        "il prato era pieno — e stamattina era vuoto, pulito, come sempre.\n\n"
        "grazie.\n\n"
        "ci vediamo al prossimo mercoledì.\n"
        "#illenznuolobianco"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(post_evento)
    set_font(r, "Georgia", 10, italic=True, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(8)

    # ── 6. Piano editoriale ──
    add_heading(doc, "6. PIANO EDITORIALE — 3 SETTIMANE ALL'EVENTO", 1)
    add_body(doc, "Ipotesi: evento mercoledì 11 giugno 2026", italic=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Data", "Piattaforma", "Contenuto", "Obiettivo"]
    for i, (cell, h) in enumerate(zip(hdr, headers)):
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                run.bold = True
        shade_cell(cell, "D0D0D0")

    for row_data in PIANO_EDITORIALE:
        add_table_row(table, row_data, bold_first=True)

    # ── 7. Meccanismo di crescita ──
    add_heading(doc, "7. MECCANISMO DI CRESCITA", 1)

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    for i, (cell, h) in enumerate(zip(hdr2, ["Edizione", "Obiettivo", "Leve"])):
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                run.bold = True
        shade_cell(cell, "D0D0D0")

    for row_data in CRESCITA:
        add_table_row(table2, row_data, bold_first=True)

    # ── 8. Cosa NON fare ──
    add_heading(doc, "8. COSA NON FARE", 1)
    non_fare = [
        "Non monetizzare subito — appena metti un biglietto, perdi l'anima del concept",
        "Non over-organizzare — la spontaneità è il valore, non smontarla con troppa struttura",
        "Non fare comunicazione troppo «event agency» — il tono deve restare personale",
        "Non fare l'evento ogni settimana — la rarità è potere",
    ]
    for n in non_fare:
        add_bullet(doc, n)

    # ── 9. Il sogno lungo ──
    add_heading(doc, "9. IL SOGNO LUNGO (anno 2+)", 1)
    add_body(doc, (
        "Come la Cena in Bianco parigina è stata replicata in tutto il mondo, "
        "Il Lenzuolo Bianco può diventare un appuntamento che altre città chiedono di replicare — "
        "con Zeno Govoni come fondatore del format, esattamente come è avvenuto con Street Dinner."
    ))

    # ── Note finale ──
    add_heading(doc, "10. NOTA FINALE", 1)
    add_body(doc, (
        "Il contenuto più potente non lo farà l'organizzatore — lo faranno i partecipanti. "
        "Il lavoro nei giorni prima è seminare bene il concept. "
        "Il lavoro nei giorni dopo è raccogliere e amplificare quello che hanno creato gli altri. "
        "Reposta, ringrazia, nomina. Ogni lenzuolo condiviso è pubblicità gratuita e autentica."
    ))

    doc.save(OUTPUT_DOCX)
    print(f"✓ DOCX salvato: {OUTPUT_DOCX}")


# ─── PDF (ReportLab) ──────────────────────────────────────────────────────────

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

W, H = A4

def build_pdf():
    doc = SimpleDocTemplate(
        OUTPUT_PDF,
        pagesize=A4,
        leftMargin=3*cm, rightMargin=3*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )

    styles = getSampleStyleSheet()

    def sty(name, parent="Normal", **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    cover_title = sty("CoverTitle", fontSize=28, leading=34,
                      alignment=TA_CENTER, textColor=colors.HexColor("#141414"),
                      fontName="Times-Bold", spaceAfter=10)
    cover_sub   = sty("CoverSub", fontSize=13, leading=18,
                      alignment=TA_CENTER, textColor=colors.HexColor("#646464"),
                      fontName="Times-Italic", spaceAfter=6)
    cover_meta  = sty("CoverMeta", fontSize=10, leading=14,
                      alignment=TA_CENTER, textColor=colors.HexColor("#8C8C8C"),
                      fontName="Times-Roman")

    h1 = sty("H1", fontSize=17, leading=22, fontName="Times-Bold",
             textColor=colors.HexColor("#141414"),
             spaceBefore=20, spaceAfter=8)
    h2 = sty("H2", fontSize=13, leading=18, fontName="Times-Bold",
             textColor=colors.HexColor("#323232"),
             spaceBefore=14, spaceAfter=6)
    body = sty("Body", fontSize=11, leading=16, fontName="Times-Roman",
               textColor=colors.HexColor("#3C3C3C"),
               spaceAfter=6, alignment=TA_JUSTIFY)
    italic_body = sty("ItalicBody", fontSize=11, leading=16,
                      fontName="Times-Italic",
                      textColor=colors.HexColor("#505050"),
                      spaceAfter=6, alignment=TA_JUSTIFY)
    bullet_sty = sty("Bullet", fontSize=11, leading=16, fontName="Times-Roman",
                     textColor=colors.HexColor("#3C3C3C"),
                     leftIndent=20, spaceAfter=4,
                     bulletIndent=10)
    manifesto_title = sty("ManifTitle", fontSize=14, leading=20,
                          fontName="Times-Bold",
                          textColor=colors.HexColor("#141414"),
                          leftIndent=40, spaceAfter=4)
    manifesto_line = sty("ManifLine", fontSize=11, leading=16,
                         fontName="Times-Italic",
                         textColor=colors.HexColor("#323232"),
                         leftIndent=40, spaceAfter=3)
    manifesto_firma = sty("ManifFirma", fontSize=10, leading=14,
                          fontName="Times-Italic",
                          textColor=colors.HexColor("#787878"),
                          leftIndent=40, spaceAfter=4)
    mono = sty("Mono", fontSize=10, leading=14, fontName="Courier",
               textColor=colors.HexColor("#323232"),
               leftIndent=30, spaceAfter=3)
    quote = sty("Quote", fontSize=10, leading=15,
                fontName="Times-Italic",
                textColor=colors.HexColor("#505050"),
                leftIndent=30, spaceAfter=8,
                alignment=TA_LEFT)

    story = []

    # COPERTINA
    story += [
        Spacer(1, 4*cm),
        Paragraph("IL LENZUOLO BIANCO", cover_title),
        Spacer(1, 0.4*cm),
        Paragraph("Piano evento virale — Ferrara, giugno 2026", cover_sub),
        Spacer(1, 0.2*cm),
        Paragraph(f"Ideato da Zeno Govoni  •  {datetime.date.today().strftime('%d/%m/%Y')}", cover_meta),
        PageBreak(),
    ]

    def h(text, level=1):
        s = h1 if level == 1 else h2
        story.append(Paragraph(text, s))

    def b(text):
        story.append(Paragraph(text, body))

    def bi(text):
        story.append(Paragraph(text, italic_body))

    def bullet(text):
        story.append(Paragraph(f"• {text}", bullet_sty))

    def sp(n=6):
        story.append(Spacer(1, n))

    def hr():
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#C0C0C0"), spaceAfter=8))

    # 1. IL CONCEPT
    h("1. IL CONCEPT")
    hr()
    b("Il Lenzuolo Bianco è un evento flash mob ad ingresso libero, che si tiene il mercoledì "
      "sera all'imbrunire, a giugno, nel sottomura di Ferrara. I partecipanti portano un grande "
      "lenzuolo bianco — meglio se di canapa, trovato in un mercatino — si stendono sul prato "
      "esterno alle mura, portando cibo, bevande fredde e la propria musica. L'evento si conclude "
      "lasciando il luogo esattamente come era all'arrivo.")
    h("Perché può diventare virale", 2)
    for bu in [
        "Visivo e fotografabile — lenzuola bianche sul verde, mura medievali, luce del tramonto",
        "Bassa barriera d'ingresso — nessun biglietto, nessuna prenotazione, nessun dress code",
        "Altamente condivisibile — ogni lenzuolo è un protagonista diverso",
        "Nostalgia + creatività — il lenzuolo di canapa è un oggetto con anima",
        "La regola del rispetto — «lascia il posto come lo hai trovato» è un valore",
    ]:
        bullet(bu)
    sp()

    # 2. PRIMA EDIZIONE
    h("2. PRIMA EDIZIONE — PICCOLA E CURATA")
    hr()
    bi("Non fare grande subito. Fai bello.")
    for s in [
        "Invita 30–50 persone che conosci e che sai che postano bene",
        "Scegli un mercoledì di giugno con meteo affidabile (metà mese)",
        "Scatta foto e video con un fotografo amico, non uno studio pagato",
        "Crea contenuti prima, durante e dopo l'evento",
        "L'obiettivo del primo evento non è la folla — è il materiale visivo per le edizioni successive",
    ]:
        bullet(s)
    sp()

    # 3. IDENTITÀ VISIVA
    h("3. IDENTITÀ VISIVA E NOME")
    hr()
    b("Il nome «Il Lenzuolo Bianco» è evocativo, italiano, ironico, caldo — non va cambiato. "
      "Va costruita un'identità visiva coerente attorno ad esso.")
    for i in [
        "Logo: minimalista — un lenzuolo che sventola o una sagoma stilizzata",
        "Palette: bianco / verde prato / mattone rosso delle mura ferraresi",
        "Hashtag: #illenznuolobianco — uno solo, usato sempre",
        "Profilo Instagram dedicato — separato da Street Dinner e Hotel Annunziata",
    ]:
        bullet(i)
    sp()

    # 4. IL MANIFESTO
    h("4. IL MANIFESTO")
    hr()
    bi("Da stampare, condividere, incorniciare.")
    sp(8)
    for line in MANIFESTO:
        if line == "Il Lenzuolo Bianco":
            story.append(Paragraph(line, manifesto_title))
        elif line.startswith("—"):
            story.append(Paragraph(line, manifesto_firma))
        elif line == "":
            sp(4)
        else:
            story.append(Paragraph(line, manifesto_line))
    sp()

    # 5. COMUNICAZIONE SOCIAL
    h("5. COMUNICAZIONE SOCIAL")
    hr()
    h("Tono di voce", 2)
    b("Non è un evento. È un'idea. Il tono è sempre personale, mai da «agenzia eventi». "
      "Breve, poetico, con un pizzico di ironia. Niente maiuscole urlate, niente emoji a raffica.")

    h("Bio Instagram suggerita", 2)
    for line in ["il lenzuolo bianco", "un prato. le mura. un mercoledì di giugno.",
                 "📍 Ferrara — sottomura", "#illenznuolobianco"]:
        story.append(Paragraph(line, mono))
    sp()

    h("Fase 1 — Teaser (3 settimane prima)", 2)
    for title, caption in [
        ("Post 1 — foto mura al tramonto",
         "c'è un'idea che gira da un po'.\nha bisogno di un prato, di qualche lenzuolo e di una serata di giugno.\npresto."),
        ("Post 2 — mano con lenzuolo di canapa",
         "nei mercatini ne trovi ancora.\npesanti, ruvidi, bianchi.\nperfetti."),
        ("Post 3 — prato vuoto, mura sullo sfondo",
         "mercoledì 11 giugno.\nore 19.00.\nsottomura, lato est.\nporta un lenzuolo."),
    ]:
        story.append(Paragraph(f"<b>{title}</b>", ParagraphStyle("bh", parent=styles["Normal"],
            fontSize=10, fontName="Times-Bold", leftIndent=0, spaceAfter=2,
            textColor=colors.HexColor("#505050"))))
        story.append(Paragraph(caption.replace("\n", "<br/>"), quote))

    h("Fase 2 — Annuncio ufficiale (2 settimane prima)", 2)
    annuncio_pdf = (
        "<b>Il Lenzuolo Bianco</b><br/><br/>"
        "un evento senza palco, senza biglietto, senza programma.<br/><br/>"
        "porti un lenzuolo bianco — grande, meglio se vecchio.<br/>"
        "porti da bere, da mangiare, la tua musica.<br/>"
        "stenditi sul prato delle mura.<br/>"
        "conosci chi ti sta accanto.<br/><br/>"
        "mercoledì 11 giugno, ore 19.00<br/>"
        "sottomura di Ferrara — prato esterno lato Rampari<br/><br/>"
        "prima di andare, lasci il posto come l'hai trovato.<br/><br/>"
        "#illenznuolobianco"
    )
    story.append(Paragraph(annuncio_pdf, quote))

    h("Fase 3 — Hype (settimana prima)", 2)
    guida_pdf = (
        "guida minima al lenzuolo bianco:<br/><br/>"
        "✦ il lenzuolo: bianco. più vecchio meglio è. di canapa se ce l'hai.<br/>"
        "✦ da bere: freddo. buono. condivisibile.<br/>"
        "✦ da mangiare: niente piatti caldi. aperitivo lungo.<br/>"
        "✦ musica: sì, ma con rispetto del vicino di lenzuolo.<br/>"
        "✦ orario: dalle 19. finché c'è luce e voglia.<br/>"
        "✦ il prato: lascialo come lo trovi.<br/><br/>"
        "il resto è tuo."
    )
    story.append(Paragraph(guida_pdf, quote))

    h("Fase 4 — Post evento", 2)
    post_pdf = (
        "eravate in tanti.<br/>"
        "ognuno con il suo lenzuolo, il suo cibo, la sua musica.<br/>"
        "il prato era pieno — e stamattina era vuoto, pulito, come sempre.<br/><br/>"
        "grazie.<br/><br/>"
        "ci vediamo al prossimo mercoledì.<br/>"
        "#illenznuolobianco"
    )
    story.append(Paragraph(post_pdf, quote))
    sp()

    # 6. PIANO EDITORIALE
    h("6. PIANO EDITORIALE — 3 SETTIMANE ALL'EVENTO")
    hr()
    bi("Ipotesi: evento mercoledì 11 giugno 2026")
    sp(6)

    table_data = [["Data", "Piattaforma", "Contenuto", "Obiettivo"]]
    for row in PIANO_EDITORIALE:
        table_data.append(list(row))

    col_widths = [2.8*cm, 2.5*cm, 8.5*cm, 3.8*cm]
    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("FONTNAME",    (0,0), (-1,0),  "Times-Bold"),
        ("FONTNAME",    (0,1), (-1,-1), "Times-Roman"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("LEADING",     (0,0), (-1,-1), 13),
        ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#D0D0D0")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#F5F5F5")]),
        ("GRID",        (0,0), (-1,-1), 0.4, colors.HexColor("#C0C0C0")),
        ("VALIGN",      (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",  (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ("LEFTPADDING", (0,0), (-1,-1), 5),
        ("RIGHTPADDING",(0,0), (-1,-1), 5),
        ("FONTNAME",    (0,1), (0,-1),  "Times-Bold"),
    ]))
    story.append(tbl)
    sp()

    # 7. CRESCITA
    h("7. MECCANISMO DI CRESCITA")
    hr()
    table_data2 = [["Edizione", "Obiettivo", "Leve"]]
    for row in CRESCITA:
        table_data2.append(list(row))
    col_widths2 = [4.5*cm, 6.5*cm, 6.5*cm]
    tbl2 = Table(table_data2, colWidths=col_widths2, repeatRows=1)
    tbl2.setStyle(TableStyle([
        ("FONTNAME",    (0,0), (-1,0),  "Times-Bold"),
        ("FONTNAME",    (0,1), (-1,-1), "Times-Roman"),
        ("FONTSIZE",    (0,0), (-1,-1), 10),
        ("LEADING",     (0,0), (-1,-1), 14),
        ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#D0D0D0")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#F5F5F5")]),
        ("GRID",        (0,0), (-1,-1), 0.4, colors.HexColor("#C0C0C0")),
        ("VALIGN",      (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",  (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING", (0,0), (-1,-1), 5),
        ("FONTNAME",    (0,1), (0,-1),  "Times-Bold"),
    ]))
    story.append(tbl2)
    sp()

    # 8. COSA NON FARE
    h("8. COSA NON FARE")
    hr()
    for n in [
        "Non monetizzare subito — appena metti un biglietto, perdi l'anima del concept",
        "Non over-organizzare — la spontaneità è il valore, non smontarla con troppa struttura",
        "Non fare comunicazione troppo «event agency» — il tono deve restare personale",
        "Non fare l'evento ogni settimana — la rarità è potere",
    ]:
        bullet(n)
    sp()

    # 9. IL SOGNO LUNGO
    h("9. IL SOGNO LUNGO (anno 2+)")
    hr()
    b("Come la Cena in Bianco parigina è stata replicata in tutto il mondo, "
      "Il Lenzuolo Bianco può diventare un appuntamento che altre città chiedono di replicare — "
      "con Zeno Govoni come fondatore del format, esattamente come è avvenuto con Street Dinner.")
    sp()

    # 10. NOTA FINALE
    h("10. NOTA FINALE")
    hr()
    b("Il contenuto più potente non lo farà l'organizzatore — lo faranno i partecipanti. "
      "Il lavoro nei giorni prima è seminare bene il concept. "
      "Il lavoro nei giorni dopo è raccogliere e amplificare quello che hanno creato gli altri. "
      "Reposta, ringrazia, nomina. Ogni lenzuolo condiviso è pubblicità gratuita e autentica.")
    sp(20)
    story.append(Paragraph("Il Lenzuolo Bianco — Ferrara", sty("Footer",
        fontSize=9, fontName="Times-Italic",
        alignment=TA_CENTER, textColor=colors.HexColor("#AAAAAA"))))

    doc.build(story)
    print(f"✓ PDF salvato:  {OUTPUT_PDF}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("\nFatto! Entrambi i file sono sul Desktop.")
